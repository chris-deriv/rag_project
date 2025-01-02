"""Create test files for document processing tests."""
import os
from fpdf import FPDF
from docx import Document

def create_test_pdf():
    """Create a test PDF with a known title and content."""
    pdf = FPDF()
    
    # Set document info
    pdf.set_title("Test Document Title")
    
    # Add a page
    pdf.add_page()
    
    # Add title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Test Document Title", ln=True, align="C")
    
    # Add content
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 10, "This is a test document.\n\nIt contains multiple paragraphs to test the chunking functionality.\n\nEach paragraph should be processed correctly.")
    
    # Save the file
    output_path = os.path.join("tests", "test_data", "test.pdf")
    pdf.output(output_path)
    print(f"Created test PDF at: {output_path}")

def create_test_docx():
    """Create a test DOCX with a known title and content."""
    doc = Document()
    
    # Set core properties
    doc.core_properties.title = "Test DOCX Title"
    
    # Add content
    doc.add_heading("Test DOCX Title", 0)
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It contains multiple paragraphs to test the chunking functionality.")
    doc.add_paragraph("Each paragraph should be processed correctly.")
    
    # Save the file
    output_path = os.path.join("tests", "test_data", "test.docx")
    doc.save(output_path)
    print(f"Created test DOCX at: {output_path}")

if __name__ == "__main__":
    # Create test files
    create_test_pdf()
    create_test_docx()
    
    # Create test files with different names for fallback tests
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 10, "Test content")
    pdf.output(os.path.join("tests", "test_data", "test_document.pdf"))
    
    doc = Document()
    doc.add_paragraph("Test content")
    doc.save(os.path.join("tests", "test_data", "test_document.docx"))
    
    print("All test files created successfully!")
