"""Document processing for the RAG application with advanced chunking strategies."""
import os
import uuid
from typing import List, Dict, Optional, BinaryIO, Union, Any
from dataclasses import dataclass
import re
from pypdf import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import tiktoken
import logging
import warnings
import subprocess
import tempfile
from .database import VectorDatabase
from .embedding import EmbeddingGenerator

# Configure logging with immediate output
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',
    force=True
)
logger = logging.getLogger(__name__)

@dataclass
class ProcessingState:
    """Tracks the state of document processing."""
    status: str  # 'processing', 'completed', 'error'
    error: Optional[str] = None
    source_name: Optional[str] = None
    chunk_count: int = 0
    total_chunks: int = 0

@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document with metadata."""
    id: str
    text: str
    metadata: Dict

class DocumentProcessor:
    """Handles document processing with advanced chunking strategies."""
    
    def __init__(self, 
                 chunk_size: int = 500,
                 chunk_overlap: int = 50,
                 length_function: str = "char"):
        """Initialize the document processor."""
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.length_function = length_function
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        
        # Initialize text splitter with better separators
        self.text_splitter = RecursiveCharacterTextSplitter(
            separators=["\n\n", "\n", ".", "!", "?", ";", ":", " ", ""],
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=self._get_length_function(),
            is_separator_regex=False
        )

    def _get_length_function(self) -> callable:
        """Get the appropriate length function based on settings."""
        if self.length_function == "token":
            return lambda x: len(self.tokenizer.encode(x))
        return len
    
    def process_document(self, file_path: str) -> List[DocumentChunk]:
        """Process a document file into chunks with metadata."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext == '.pdf':
            sections = self._extract_pdf_text(file_path)
        elif file_ext in ['.doc', '.docx']:
            sections = self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
            
        chunks = []
        for section in sections:
            chunk = DocumentChunk(
                id=str(uuid.uuid4()),
                text=section['text'],
                metadata=section['metadata']
            )
            chunks.append(chunk)
            
        return chunks
        
    def _extract_pdf_text(self, file_path: str) -> List[Dict]:
        """Extract text and metadata from PDF file."""
        sections = []
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    sections.append({
                        'text': text,
                        'metadata': {
                            'source_name': os.path.basename(file_path),
                            'title': os.path.splitext(os.path.basename(file_path))[0],
                            'file_type': 'pdf',
                            'section_type': 'content',
                            'chunk_index': i,
                            'total_chunks': total_pages
                        }
                    })
        except Exception as e:
            raise ValueError(f"Error processing PDF: {str(e)}")
            
        return sections
        
    def _extract_docx_text(self, file_path: str) -> List[Dict]:
        """Extract text and metadata from DOCX file."""
        sections = []
        try:
            if file_path.endswith('.doc'):
                # Convert DOC to DOCX using LibreOffice
                # LibreOffice keeps original name but changes extension
                original_name = os.path.basename(file_path)
                docx_name = original_name.rsplit('.', 1)[0] + '.docx'
                docx_path = os.path.join('/app/tmp', docx_name)
                
                # Log current state
                logger.info(f"Starting conversion of {file_path}")
                logger.info(f"Expected output: {docx_path}")
                logger.info(f"Current tmp contents: {os.listdir('/app/tmp')}")
                
                # Run LibreOffice conversion
                result = subprocess.run(
                    ['soffice', '--headless', '--convert-to', 'docx', '--outdir', '/app/tmp', file_path],
                    capture_output=True,
                    text=True,
                    check=True
                )
                
                # Log conversion results
                logger.info(f"LibreOffice stdout: {result.stdout}")
                if result.stderr:
                    logger.warning(f"LibreOffice stderr: {result.stderr}")
                logger.info(f"Tmp contents after conversion: {os.listdir('/app/tmp')}")
                
                # Verify conversion
                if not os.path.exists(docx_path):
                    raise ValueError(
                        f"LibreOffice conversion failed. Expected file not found at {docx_path}. "
                        f"Directory contents: {os.listdir('/app/tmp')}. "
                        f"Command output: {result.stdout}. "
                        f"Error output: {result.stderr}"
                    )
                
                file_path = docx_path
            
            doc = Document(file_path)
            total_paras = len(doc.paragraphs)
            
            # Only include non-empty paragraphs and normalize indices
            non_empty_sections = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    non_empty_sections.append(text)
            
            # Create sections with normalized indices
            total_sections = len(non_empty_sections)
            for i, text in enumerate(non_empty_sections):
                sections.append({
                    'text': text,
                    'metadata': {
                        'source_name': os.path.basename(file_path),
                        'title': os.path.splitext(os.path.basename(file_path))[0],
                        'file_type': 'docx',
                        'section_type': 'content',
                        'chunk_index': i,
                        'total_chunks': total_sections
                    }
                })
        except Exception as e:
            raise ValueError(f"Error processing DOCX: {str(e)}")
            
        return sections

class DocumentStore:
    """Manages document storage and retrieval with atomic operations."""
    
    def __init__(self):
        self.processor = DocumentProcessor()
        self.db = VectorDatabase()
        self.embedding_generator = EmbeddingGenerator()
        self._processing_states = {}  # Track processing states
    
    def get_processing_state(self, filename: str) -> Optional[ProcessingState]:
        """Get the current processing state for a document."""
        return self._processing_states.get(filename)
    
    def _update_processing_state(self, filename: str, state: ProcessingState) -> None:
        """Update the processing state for a document."""
        self._processing_states[filename] = state
        logger.info(f"Updated processing state for {filename}: {state}")
    
    def process_and_store_document(self, file_path: str) -> ProcessingState:
        """
        Process and store a document with atomic operations.
        This is the single point of entry for document processing.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ProcessingState: Final state of document processing
        """
        filename = os.path.basename(file_path)
        state = ProcessingState(status='processing')
        self._update_processing_state(filename, state)
        
        try:
            # 1. Process document into chunks
            logger.info(f"Processing document: {filename}")
            chunks = self.processor.process_document(file_path)
            if not chunks:
                raise ValueError("No chunks generated from document")
            
            # Update state with chunk information
            state.chunk_count = len(chunks)
            state.total_chunks = len(chunks)
            state.source_name = chunks[0].metadata.get('source_name', filename)
            self._update_processing_state(filename, state)
            
            # 2. Generate embeddings (single point of embedding generation)
            logger.info("Generating embeddings...")
            texts = [chunk.text for chunk in chunks]
            embeddings = self.embedding_generator.generate_embeddings(texts)
            
            # 3. Prepare documents for database
            documents = []
            for i, chunk in enumerate(chunks):
                doc = {
                    "id": chunk.id,
                    "text": chunk.text,
                    "embedding": embeddings[i],
                    **chunk.metadata
                }
                documents.append(doc)
            
            # 4. Delete any existing document with same source name
            logger.info(f"Checking for existing document: {state.source_name}")
            existing_chunks = self.db.get_document_chunks(state.source_name)
            if existing_chunks:
                logger.info(f"Found existing document with {len(existing_chunks)} chunks. Removing...")
                # Get existing IDs and delete them
                existing_ids = [chunk['id'] for chunk in existing_chunks]
                self.db.collection.delete(ids=existing_ids)
                logger.info(f"Deleted {len(existing_ids)} existing chunks")
            
            # 5. Atomic database operation
            logger.info("Adding documents to database...")
            self.db.add_documents(documents)
            
            # 6. Verify storage and chunk consistency
            stored_chunks = self.db.get_document_chunks(state.source_name)
            if not stored_chunks:
                raise ValueError(f"Storage verification failed - no chunks found for {filename}")
            
            if len(stored_chunks) != len(chunks):
                raise ValueError(
                    f"Storage verification failed - chunk count mismatch for {filename}. "
                    f"Expected {len(chunks)}, found {len(stored_chunks)}"
                )
            
            # Verify chunk indices and total_chunks are consistent
            chunk_indices = sorted(int(chunk.get('chunk_index', -1)) for chunk in stored_chunks)
            expected_indices = list(range(len(chunks)))
            if chunk_indices != expected_indices:
                raise ValueError(
                    f"Storage verification failed - inconsistent chunk indices for {filename}. "
                    f"Expected sequential indices 0-{len(chunks)-1}, got {chunk_indices}"
                )
            
            # Verify total_chunks matches actual count
            for chunk in stored_chunks:
                if int(chunk.get('total_chunks', 0)) != len(chunks):
                    raise ValueError(
                        f"Storage verification failed - total_chunks mismatch for {filename}. "
                        f"Expected {len(chunks)}, got {chunk.get('total_chunks')}"
                    )
            
            # Update final state
            state.status = 'completed'
            self._update_processing_state(filename, state)
            logger.info(f"Successfully processed and stored {filename}")
            
            return state
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error processing document {filename}: {error_msg}")
            state.status = 'error'
            state.error = error_msg
            self._update_processing_state(filename, state)
            raise
    
    def get_documents(self) -> List[Dict[str, str]]:
        """Return all document chunks."""
        return self.db.get_all_documents()
    
    def get_document_info(self, source_name: str) -> Optional[Dict[str, Any]]:
        """Get information about a specific document."""
        chunks = self.db.get_document_chunks(source_name)
        if not chunks:
            return None
            
        return {
            'source_name': source_name,
            'title': chunks[0].get('title', ''),
            'chunk_count': len(chunks),
            'total_chunks': chunks[0].get('total_chunks', len(chunks))
        }

# Initialize global document store
document_store = DocumentStore()

# Expose simplified API
def process_document(file_path: str) -> List[Dict[str, str]]:
    """Process a document and return its chunks."""
    state = document_store.process_and_store_document(file_path)
    if state.status == 'error':
        raise ValueError(f"Document processing failed: {state.error}")
    return document_store.db.get_document_chunks(state.source_name)

def get_documents() -> List[Dict[str, str]]:
    """Get all documents from the store."""
    return document_store.get_documents()

def get_processing_state(filename: str) -> Optional[ProcessingState]:
    """Get processing state for a document."""
    return document_store.get_processing_state(filename)
